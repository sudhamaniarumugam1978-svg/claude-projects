#!/usr/bin/env python3
"""Render the BuildWise Business Plan (report_content.BLOCKS) to a professional
Word document using python-docx.

Uses native Times New Roman (renders with the real font and the Rupee glyph in
Microsoft Word / Google Docs).  Page breaks are inserted at every chapter and
pagebreak block so the .docx mirrors the 30-page structure of the PDF.  Output
overwrites the existing ``BuildWise_Business_Plan_Report.docx``.
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

import report_content as RC

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "BuildWise_Business_Plan_Report.docx")

NAVY = "163A70"
NAVY_MD = "2E5A9E"
STEEL = "4E77B0"
GREY = "54606E"
LGREY = "8A94A0"
ROW_ALT = "EEF3FA"
CARD_BG = "F2F5FA"
CO_BG = "EEF1F6"
WHITE = "FFFFFF"
BLACK = "1A1A1A"

FONT = "Times New Roman"
CW_IN = 6.27          # usable content width in inches (A4, 1in margins)
_figcount = [0]


# ---------------------------------------------------------------------------
# low-level xml helpers
# ---------------------------------------------------------------------------
def _shade(el, hex_fill):
    tcPr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def cell_margins(cell, top=40, bottom=40, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, v in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def thin_borders(table, color="BFC7D2", sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def cell_left_border(cell, color=NAVY, sz=24):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)
    borders.append(e)
    tcPr.append(borders)


def set_table_width(table, width_in):
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_in * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


INLINE = re.compile(r"(\*\*.+?\*\*)")


def add_runs(par, text, size, color=BLACK, bold=False, italic=False, font=FONT):
    for part in INLINE.split(text):
        if not part:
            continue
        b = bold
        t = part
        if part.startswith("**") and part.endswith("**"):
            b = True
            t = part[2:-2]
        r = par.add_run(t)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold = b
        r.font.italic = italic
        r.font.color.rgb = RGBColor.from_string(color)
        rf = r._element.get_or_add_rPr().get_or_add_rFonts()
        rf.set(qn("w:eastAsia"), font)


# ---------------------------------------------------------------------------
# document + styles
# ---------------------------------------------------------------------------
doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.different_first_page_header_footer = True

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)
rf = normal.element.get_or_add_rPr().get_or_add_rFonts()
rf.set(qn("w:eastAsia"), FONT)


def para(space_after=6, space_before=0, align=None, line=1.15, keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    if keep_next:
        pf.keep_with_next = True
    return p


def bottom_border(p, color=NAVY, sz=12, space=2):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    pbdr.append(b)
    pPr.append(pbdr)


# ---------------------------------------------------------------------------
# block renderers
# ---------------------------------------------------------------------------
def r_chapter(label, title):
    doc.add_page_break()
    p = para(space_after=0, keep_next=True)
    add_runs(p, label, 12.5, color=STEEL, bold=True)
    p2 = para(space_after=2, keep_next=True)
    add_runs(p2, title, 20, color=NAVY, bold=True)
    bottom_border(p2, NAVY, 16)


def r_fmhead(text, rule):
    p = para(space_after=6, space_before=2, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
    add_runs(p, text, 18, color=NAVY, bold=True)
    if rule:
        bottom_border(p, NAVY, 14)


def r_heading(text, lvl):
    size = 13.5 if lvl == 2 else 12
    p = para(space_after=3, space_before=8 if lvl == 2 else 6, keep_next=True)
    add_runs(p, text, size, color=NAVY, bold=True)


def r_para(text):
    p = para(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_runs(p, text, 11)


def r_bullets(items, title, check):
    if title:
        p = para(space_after=2, space_before=4, keep_next=True)
        add_runs(p, title, 12, color=NAVY, bold=True)
    mark = "\u2714" if check else "\u2022"
    for it in items:
        p = para(space_after=2)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(mark + "  ")
        r.font.name = FONT
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r.font.bold = True
        add_runs(p, it, 11)


def r_table(header, rows, widths):
    n = len(header)
    if widths is None:
        widths = [1.0 / n] * n
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, CW_IN)
    thin_borders(t)
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr[i].width = Inches(widths[i] * CW_IN)
        cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        add_runs(p, htext, 10, color=WHITE, bold=True)
        shade_cell(hdr[i], NAVY)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i in range(n):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[i].width = Inches(widths[i] * CW_IN)
            cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            add_runs(p, str(row[i]) if i < len(row) else "", 10)
            if ri % 2 == 1:
                shade_cell(cells[i], ROW_ALT)
    para(space_after=4, space_before=0)


def r_kpis(cardlist):
    n = len(cardlist)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, CW_IN)
    no_borders(t)
    for i, (big, small) in enumerate(cardlist):
        c = t.rows[0].cells[i]
        c.width = Inches(CW_IN / n)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_margins(c, top=110, bottom=110, left=80, right=80)
        shade_cell(c, NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        add_runs(p, big, 14, color=WHITE, bold=True)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(1)
        add_runs(p2, small, 8.5, color="D6E1F0")
    _spacer_between(t)
    para(space_after=5)


def _spacer_between(t):
    # add a hair of spacing between kpi/card columns via cell margins already
    pass


def r_cards(items, cols):
    i = 0
    while i < len(items):
        chunk = items[i:i + cols]
        t = doc.add_table(rows=1, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(t, CW_IN)
        no_borders(t)
        for j in range(cols):
            c = t.rows[0].cells[j]
            c.width = Inches(CW_IN / cols)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell_margins(c, top=0, bottom=0, left=50, right=50)
            if j < len(chunk):
                title, body = chunk[j]
                # title bar paragraph
                tp = c.paragraphs[0]
                tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tp.paragraph_format.space_after = Pt(2)
                tp.paragraph_format.space_before = Pt(2)
                _shade(tp._p, NAVY)
                add_runs(tp, title, 10, color=WHITE, bold=True)
                shade_cell(c, CARD_BG)
                if isinstance(body, list):
                    for it in body:
                        bp = c.add_paragraph()
                        bp.paragraph_format.space_after = Pt(1)
                        bp.paragraph_format.left_indent = Inches(0.08)
                        add_runs(bp, "\u2022 " + it, 9, color="33414F")
                else:
                    bp = c.add_paragraph()
                    bp.paragraph_format.space_after = Pt(3)
                    bp.paragraph_format.space_before = Pt(2)
                    add_runs(bp, body, 9, color="33414F")
        i += cols
        para(space_after=3)


def r_callout(title, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t, CW_IN)
    no_borders(t)
    c = t.rows[0].cells[0]
    cell_margins(c, top=80, bottom=80, left=150, right=120)
    shade_cell(c, CO_BG)
    cell_left_border(c, NAVY, 30)
    tp = c.paragraphs[0]
    tp.paragraph_format.space_after = Pt(2)
    add_runs(tp, title, 11, color=NAVY, bold=True, italic=True)
    bp = c.add_paragraph()
    bp.paragraph_format.space_after = Pt(0)
    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(bp, text, 10.5, color="2A2A2A")
    para(space_after=5)


def r_figure(src, caption, w):
    path = os.path.join(BASE, src)
    iw, ih = Image.open(path).size
    disp_w = w * CW_IN
    disp_h = disp_w * ih / iw
    max_h = 3.3
    if disp_h > max_h:
        disp_h = max_h
        disp_w = disp_h * iw / ih
    p = para(space_after=2, space_before=2, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
    run = p.add_run()
    run.add_picture(path, width=Inches(disp_w))
    _figcount[0] += 1
    cp = para(space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(cp, f"Figure {_figcount[0]}. {caption}", 9.5, color=GREY, italic=True)


def r_note(text):
    para(space_after=2, space_before=10)
    p = para(align=WD_ALIGN_PARAGRAPH.CENTER)
    add_runs(p, text, 10, color=GREY, italic=True)


def r_hr():
    p = para(space_after=8, space_before=6)
    bottom_border(p, "BFC7D2", 6)


def r_refs(items):
    for i, it in enumerate(items, 1):
        p = para(space_after=2)
        p.paragraph_format.left_indent = Inches(0.25)
        add_runs(p, f"{i}.  {it}", 11)


# ---------------------------------------------------------------------------
# cover page
# ---------------------------------------------------------------------------
def build_cover():
    def C(text, size, bold=False, italic=False, color=NAVY, sa=6, sb=0):
        p = para(space_after=sa, space_before=sb, align=WD_ALIGN_PARAGRAPH.CENTER,
                 line=1.1)
        add_runs(p, text, size, color=color, bold=bold, italic=italic)
        return p

    def navline(sb=6, sa=8):
        p = para(space_after=sa, space_before=sb, align=WD_ALIGN_PARAGRAPH.CENTER)
        # 80% width navy line via a bordered single-cell table
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(t, CW_IN * 0.8)
        no_borders(t)
        c = t.rows[0].cells[0]
        cell_margins(c, top=0, bottom=0, left=0, right=0)
        cp = c.paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        bottom_border(cp, NAVY, 16)

    para(space_after=18)
    C("BUILDWISE CONSTRUCTIONS PVT. LTD.", 24, bold=True, color=NAVY, sa=8)
    C('"Building Trust Through Technology and Quality Construction"', 14,
      italic=True, color=GREY, sa=6)
    para(space_after=14)
    C("BUSINESS PLAN REPORT", 20, bold=True, color=NAVY, sa=4)
    navline(sb=4, sa=10)

    info = [("Industry", "Construction"),
            ("Business Model", "Technology-Enabled Construction Startup"),
            ("Headquarters", "Chennai, Tamil Nadu"),
            ("Initial Capital", "\u20b91.00 Crore"),
            ("Academic Year", "2026\u20132027")]
    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.allow_autofit = False
    no_borders(t)
    for ri, (k, v) in enumerate(info):
        c0, c1 = t.rows[ri].cells
        c0.width = Inches(1.9)
        c1.width = Inches(3.7)
        cell_margins(c0, top=20, bottom=20, left=0, right=0)
        cell_margins(c1, top=20, bottom=20, left=0, right=0)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_after = Pt(2)
        add_runs(p0, k, 13, color=BLACK, bold=True)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_after = Pt(2)
        add_runs(p1, ":  " + v, 13, color=BLACK)
    navline(sb=12, sa=10)

    C("PREPARED BY", 15, bold=True, color=NAVY, sa=8)
    for name, rno in [("THILAK KUMAR K", "RA2311042040014"),
                      ("ARJUN S", "RA2311042040018"),
                      ("SANJAY KUMAR A M", "RA2311042040024")]:
        C(name, 15, bold=True, color=BLACK, sa=0)
        C(rno, 12, color=GREY, sa=8)
    para(space_after=10)
    C("Entrepreneurship and Family Business Management", 13, bold=True, color=NAVY, sa=6)
    C("Department of Computer Science and Engineering (Emerging Technologies)", 13,
      color=BLACK, sa=0)
    C("& Computer Science and Business Systems (Final Year)", 13, color=BLACK, sa=8)
    C("SRM Institute of Science and Technology", 13, bold=True, color=NAVY, sa=0)
    C("Vadapalani Campus, Chennai", 13, color=BLACK, sa=0)


# ---------------------------------------------------------------------------
# footer with page numbers (body section restarts numbering at 1)
# ---------------------------------------------------------------------------
def add_page_field(paragraph):
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(GREY)
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(e)


def setup_body_footer(section, start=1):
    # restart page numbering at `start` for this section
    sectPr = section._sectPr
    pgNum = OxmlElement("w:pgNumType")
    pgNum.set(qn("w:start"), str(start))
    sectPr.append(pgNum)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Page ")
    r1.font.name = FONT
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor.from_string(GREY)
    add_page_field(fp)


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------
def build():
    _figcount[0] = 0
    body_started = [False]

    def start_body_section():
        # insert a section break so cover is its own (unnumbered) section
        new = doc.add_section(WD_SECTION.NEW_PAGE)
        new.page_height = Cm(29.7)
        new.page_width = Cm(21.0)
        new.top_margin = Inches(1)
        new.bottom_margin = Inches(1)
        new.left_margin = Inches(1)
        new.right_margin = Inches(1)
        new.different_first_page_header_footer = False
        setup_body_footer(new, start=1)
        # running header (right aligned) with rule
        hdr = new.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_runs(hp, "BuildWise Constructions Pvt. Ltd.", 9, color=LGREY, italic=True)
        bottom_border(hp, "BFC7D2", 6)

    for blk in RC.BLOCKS:
        t = blk["t"]
        if t == "cover":
            build_cover()
        elif t == "pagebreak":
            if not body_started[0]:
                start_body_section()
                body_started[0] = True
            else:
                doc.add_page_break()
        elif t == "chapter":
            if not body_started[0]:
                start_body_section()
                body_started[0] = True
                # first chapter: no extra page break (section already broke)
                p = para(space_after=0, keep_next=True)
                add_runs(p, blk["label"], 12.5, color=STEEL, bold=True)
                p2 = para(space_after=2, keep_next=True)
                add_runs(p2, blk["title"], 20, color=NAVY, bold=True)
                bottom_border(p2, NAVY, 16)
            else:
                r_chapter(blk["label"], blk["title"])
        elif t == "fmhead":
            r_fmhead(blk["text"], blk.get("rule", False))
        elif t == "h":
            r_heading(blk["text"], blk["lvl"])
        elif t == "p":
            r_para(blk["text"])
        elif t == "bullets":
            r_bullets(blk["items"], blk.get("title"), blk.get("check", True))
        elif t == "table":
            r_table(blk["header"], blk["rows"], blk["widths"])
        elif t == "kpis":
            r_kpis(blk["cards"])
        elif t == "cards":
            r_cards(blk["items"], blk["cols"])
        elif t == "callout":
            r_callout(blk["title"], blk["text"])
        elif t == "figure":
            r_figure(blk["src"], blk["caption"], blk["w"])
        elif t == "note":
            r_note(blk["text"])
        elif t == "hr":
            r_hr()
        elif t == "refs":
            r_refs(blk["items"])

    doc.save(OUT)
    print("Saved:", OUT, "| figures embedded:", _figcount[0])


if __name__ == "__main__":
    build()
