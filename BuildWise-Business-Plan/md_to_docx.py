#!/usr/bin/env python3
"""Convert the BuildWise Business Plan Markdown report into a professionally
formatted .docx following a single consistent typography design system.

Design system
-------------
  Headings  : Cambria Bold, Dark Navy Blue (#163A6B)
  Body      : Times New Roman 11.5 pt, justified, 1.15 spacing
  Tables    : Cambria Bold 11 pt header (white on navy) / Times New Roman 10.5 pt
  Header/Footer : Times New Roman 9 pt
"""

import os
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "Business_Plan_Report.md")
OUT = os.path.join(BASE, "BuildWise_Business_Plan_Report.docx")

NAVY = "163A6B"
GREY = "595959"
BLACK = "000000"
ROW_ALT = "EEF2F8"

HEAD_FONT = "Cambria"
BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"

FRONT_MATTER = {"DECLARATION", "ACKNOWLEDGEMENT", "TABLE OF CONTENTS"}

doc = Document()

# ---------------- Page geometry ----------------
sec = doc.sections[0]
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(2.4)
sec.different_first_page_header_footer = True  # cover has no header/footer

# ---------------- Base styles ----------------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor.from_string(BLACK)
pf = normal.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(8)
pf.line_spacing = 1.3
pf.widow_control = True
# ensure east-asian/cs also map to Times
rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
rpr.set(qn("w:eastAsia"), BODY_FONT)

# Load chapter start pages (produced by md_to_pdf.py) to fill TOC page numbers.
CHAP_PAGES = []
_cp = os.path.join(BASE, "chapter_pages.json")
if os.path.exists(_cp):
    try:
        CHAP_PAGES = json.load(open(_cp))
    except Exception:
        CHAP_PAGES = []

FIG = {"n": 0}


def fig_caption(caption):
    FIG["n"] += 1
    mm = re.match(r"^Figure\s*\d*\s*:\s*(.*)$", caption)
    desc = mm.group(1) if mm else caption
    return f"Figure {FIG['n']}: {desc}"


for lvl, sz in [(1, 18), (2, 15), (3, 13), (4, 12)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = HEAD_FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(NAVY)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.space_before = Pt(10 if lvl <= 2 else 6)
    st.paragraph_format.space_after = Pt(6 if lvl == 1 else 4)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), HEAD_FONT)


def _set_run(run, font=BODY_FONT, size=12, bold=False, italic=False, color=BLACK):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rf = run._element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), font)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+\*)")


def add_runs(paragraph, text, font=BODY_FONT, size=12, color=BLACK, base_bold=False):
    """Add text honouring **bold** and *italic* inline markup."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            _set_run(r, font, size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            _set_run(r, font, size, bold=base_bold, italic=True, color=color)
        else:
            r = paragraph.add_run(part)
            _set_run(r, font, size, bold=base_bold, color=color)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(table, top=60, bottom=60, left=110, right=110):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                     ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def add_table(rows):
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = []
    for r in rows[2:]:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        body.append(cells)
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(table)

    hdr = table.rows[0].cells
    repeat_header(table.rows[0])
    for i, h in enumerate(header):
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        add_runs(p, h.replace("**", ""), font=HEAD_FONT, size=11, color="FFFFFF")
        for run in p.runs:
            run.font.bold = True
        shade_cell(hdr[i], NAVY)

    for ridx, cells in enumerate(body):
        rowc = table.add_row().cells
        for i in range(ncol):
            txt = cells[i] if i < len(cells) else ""
            rowc[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = rowc[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_runs(p, txt, font=BODY_FONT, size=11, color=BLACK)
            if ridx % 2 == 1:
                shade_cell(rowc[i], ROW_ALT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(codelines):
    # Trim leading/trailing blank lines
    while codelines and not codelines[0].strip():
        codelines.pop(0)
    while codelines and not codelines[-1].strip():
        codelines.pop()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    for idx, ln in enumerate(codelines):
        run = p.add_run(ln if ln else " ")
        _set_run(run, font=MONO_FONT, size=9.5, color=NAVY)
        if idx < len(codelines) - 1:
            run.add_break()


def add_image(path, caption=None):
    from PIL import Image
    iw, ih = Image.open(path).size
    # Wide diagrams get near-full width; tall diagrams are capped by height.
    max_w_in = 6.4
    max_h_in = 7.4
    w_in = max_w_in
    if (w_in * ih / iw) > max_h_in:
        w_in = max_h_in * iw / ih
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(path, width=Inches(w_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        add_runs(cp, fig_caption(caption), font=BODY_FONT, size=11, color=GREY)
        for r in cp.runs:
            r.font.italic = True


def build_toc_table(items):
    """Render the Table of Contents as a bordered S.No / Contents / Page No table."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(table, top=70, bottom=70, left=130, right=130)
    headers = ["S. No", "Contents", "Page No"]
    aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
    widths = [Cm(2.2), Cm(11.6), Cm(2.6)]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr[i].width = widths[i]
        p = hdr[i].paragraphs[0]
        p.alignment = aligns[i]
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(3)
        add_runs(p, h, font=HEAD_FONT, size=12, color="FFFFFF")
        for r in p.runs:
            r.font.bold = True
        shade_cell(hdr[i], NAVY)
    for idx, title in enumerate(items, start=1):
        cells = table.add_row().cells
        page = ""
        if idx - 1 < len(CHAP_PAGES) and CHAP_PAGES[idx - 1]:
            page = str(CHAP_PAGES[idx - 1])
        vals = [str(idx), title, page]
        for i in range(3):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[i].width = widths[i]
            p = cells[i].paragraphs[0]
            p.alignment = aligns[i]
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(3)
            add_runs(p, vals[i], font=BODY_FONT, size=12, color=BLACK)
            if idx % 2 == 0:
                shade_cell(cells[i], ROW_ALT)


# ==================== COVER PAGE ====================
def centered(text, font, size, bold=False, italic=False, color=NAVY,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    _set_run(r, font, size, bold=bold, italic=italic, color=color)
    return p


def divider(space_before=8, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY)
    pbdr.append(bottom)
    pPr.append(pbdr)


def build_cover():
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    centered("BUILDWISE CONSTRUCTIONS PVT. LTD.", HEAD_FONT, 26, bold=True,
             color=NAVY, space_before=6, space_after=6)
    centered('"Building Trust Through Technology and Quality Construction"',
             HEAD_FONT, 15, italic=True, color=GREY, space_after=12)
    centered("BUSINESS PLAN REPORT", HEAD_FONT, 19, bold=True, color=NAVY,
             space_after=4)
    divider(space_before=4, space_after=10)

    # Business details -- centred table, labels bold navy, values TNR black,
    # colons perfectly aligned by the fixed column boundary.
    details = [
        ("Industry", "Construction"),
        ("Business Model", "Technology-Enabled Construction Startup"),
        ("Location", "Chennai, Tamil Nadu, India"),
        ("Initial Capital", "\u20b91.00 Crore"),
        ("Academic Year", "2026"),
    ]
    t = doc.add_table(rows=len(details), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.allow_autofit = False
    for ridx, (label, value) in enumerate(details):
        c0, c1 = t.rows[ridx].cells
        c0.width = Cm(4.6)
        c1.width = Cm(8.6)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(label)
        _set_run(r0, HEAD_FONT, 12, bold=True, color=NAVY)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(":  ")
        _set_run(r1, HEAD_FONT, 12, bold=True, color=NAVY)
        r2 = p1.add_run(value)
        _set_run(r2, BODY_FONT, 12, color=BLACK)

    divider(space_before=12, space_after=10)

    centered("PREPARED BY", HEAD_FONT, 14, bold=True, color=NAVY, space_after=6)
    team = [
        ("THILAK KUMAR K", "RA2311042040014"),
        ("ARJUN S", "RA2311042040018"),
        ("SANJAY KUMAR A M", "RA2311042040024"),
    ]
    for name, reg in team:
        centered(name, HEAD_FONT, 13, bold=True, color=BLACK, space_after=0)
        centered(reg, BODY_FONT, 11, color=GREY, space_after=5)

    divider(space_before=8, space_after=8)

    centered("Entrepreneurship and Family Business Management", HEAD_FONT, 12,
             color=NAVY, space_after=5)
    centered("Department of Computer Science and Engineering (Emerging Technologies)",
             HEAD_FONT, 12, color=BLACK, space_after=0)
    centered("& Computer Science and Business Systems (Final Year)",
             HEAD_FONT, 12, color=BLACK, space_after=7)
    centered("SRM Institute of Science and Technology", HEAD_FONT, 13, bold=True,
             color=NAVY, space_after=0)
    centered("Vadapalani Campus, Chennai", BODY_FONT, 12, color=BLACK, space_after=0)
    doc.add_page_break()   # cover occupies exactly one page


# ==================== HEADER / FOOTER ====================
def build_header_footer():
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("BuildWise Constructions Pvt. Ltd.  |  Business Plan Report")
    _set_run(hr, BODY_FONT, 9, color=GREY)
    # bottom border on header
    pPr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "BBBBBB")
    pbdr.append(bottom)
    pPr.append(pbdr)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page N field
    r1 = fp.add_run("Page ")
    _set_run(r1, BODY_FONT, 9, color=GREY)
    _add_page_field(fp)


def _add_page_field(paragraph):
    run = paragraph.add_run()
    _set_run(run, BODY_FONT, 9, color=GREY)
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldEnd)


# ==================== BUILD ====================
build_cover()
build_header_footer()

with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
first_chapter = True
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # image  ![alt](file)
    im = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", stripped)
    if im:
        img_path = os.path.join(BASE, im.group(2))
        caption = None
        # look ahead for a *Figure ...* caption line
        j = i + 1
        while j < len(lines) and lines[j].strip() == "":
            j += 1
        if j < len(lines) and re.match(r"^\*Figure.*\*$", lines[j].strip()):
            caption = lines[j].strip().strip("*")
            i = j
        if os.path.exists(img_path):
            add_image(img_path, caption)
        i += 1
        continue

    # code block
    if stripped.startswith("```"):
        block = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith("```"):
            block.append(lines[i])
            i += 1
        add_code_block(block)
        i += 1
        continue

    # table
    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1].strip()):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            tbl.append(lines[i])
            i += 1
        add_table(tbl)
        continue

    if stripped == "---":
        i += 1
        continue

    # headings
    m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip().replace("**", "")
        if level == 2 and text.upper() == "TABLE OF CONTENTS":
            # Own page + a proper bordered TOC table.
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(14)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.keep_with_next = True
            add_runs(p, "TABLE OF CONTENTS", font=HEAD_FONT, size=16, color=NAVY)
            for r in p.runs:
                r.font.bold = True
            j = i + 1
            toc_items = []
            while j < len(lines):
                s = lines[j].strip()
                if s == "" or s == "---":
                    j += 1
                    continue
                mm = re.match(r"^\d+\.\s+(.*)$", s)
                if mm:
                    toc_items.append(mm.group(1).strip())
                    j += 1
                    continue
                break
            build_toc_table(toc_items)
            i = j
            continue
        if level == 1:
            # Every chapter begins at the top of a new page.
            doc.add_page_break()
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(2)
            h.paragraph_format.keep_with_next = True
            add_runs(h, text, font=HEAD_FONT, size=18, color=NAVY)
            for r in h.runs:
                r.font.bold = True
            # navy bottom rule under chapter title
            pPr = h._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), NAVY)
            pbdr.append(bottom)
            pPr.append(pbdr)
        elif level == 2 and text.upper() in FRONT_MATTER:
            # Declaration opens its own page (cover ended with a break); the
            # short Acknowledgement flows beneath the Declaration.
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(30 if text.upper() == "ACKNOWLEDGEMENT" else 44)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(14)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.keep_with_next = True
            add_runs(p, text.upper(), font=HEAD_FONT, size=16, color=NAVY)
            for r in p.runs:
                r.font.bold = True
        else:
            h = doc.add_heading(level=min(level, 4))
            sz = {2: 15, 3: 13, 4: 12}.get(level, 12)
            add_runs(h, text, font=HEAD_FONT, size=sz, color=NAVY)
            for r in h.runs:
                r.font.bold = True
        i += 1
        continue

    # blockquote
    if stripped.startswith(">"):
        text = stripped.lstrip(">").strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_together = True
            add_runs(p, text, font=BODY_FONT, size=11, color=NAVY, base_bold=False)
            # subtle shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), ROW_ALT)
            pPr.append(shd)
        i += 1
        continue

    # bullet list
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        add_runs(p, stripped[2:])
        i += 1
        continue

    # numbered list (TOC)
    om = re.match(r"^\d+\.\s+(.*)$", stripped)
    if om:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, om.group(1))
        i += 1
        continue

    if stripped == "":
        i += 1
        continue

    # normal paragraph
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

doc.save(OUT)
print("Saved:", OUT)
